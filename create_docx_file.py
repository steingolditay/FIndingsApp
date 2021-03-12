from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches

import os
import pathlib

current_path = pathlib.Path().absolute()
dir = "docx"
font_name = "Tahoma"
title_font_size = 16
sub_title_font_size = 14
space_font_size = 14
text_font_size = 10
red = (255, 0, 0)
yellow = (255, 192, 0)
green = (112, 173, 71)






def create_docx_from_item(item):
    document = Document()

    # Title
    title_text = reformat_text_dots(item['Title'])
    title = document.add_paragraph(title_text)
    title.style = document.styles.add_style("title style", WD_STYLE_TYPE.PARAGRAPH)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    font = title.style.font
    font.name = font_name
    font.bold = True
    font.size = Pt(title_font_size)

    # Space 1
    space = document.add_paragraph()
    space.style = document.styles.add_style("space1 style ", WD_STYLE_TYPE.PARAGRAPH)
    space.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    font = space.style.font
    font.name = font_name
    font.size = Pt(space_font_size)

    # Finding Description Title
    finding_desc_title = document.add_paragraph(":פירוט הממצא")
    finding_desc_title.style = document.styles.add_style("finding description title style", WD_STYLE_TYPE.PARAGRAPH)
    finding_desc_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['finding description title style'].paragraph_format
    paragraph_format.space_after = Pt(6)
    font = finding_desc_title.style.font
    font.name = font_name
    font.size = Pt(sub_title_font_size)
    font.underline = True

    # Finding Description
    finding_desc_text = reformat_text_dots(item['Description'])
    finding_desc = document.add_paragraph(finding_desc_text)
    finding_desc.style = document.styles.add_style("finding description style", WD_STYLE_TYPE.PARAGRAPH)
    finding_desc.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['finding description style'].paragraph_format
    paragraph_format.space_after = Pt(6)
    font = finding_desc.style.font
    font.name = font_name
    font.size = Pt(text_font_size)

    # Space 2
    space = document.add_paragraph()
    space.style = document.styles.add_style("space2 style", WD_STYLE_TYPE.PARAGRAPH)
    space.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['space2 style'].paragraph_format
    paragraph_format.space_after = Pt(6)
    font = space.style.font
    font.name = font_name
    font.size = Pt(space_font_size)

    # Risk severities

    # Probability
    risk_probability = document.add_paragraph()
    risk_probability.style = document.styles.add_style("probability style", WD_STYLE_TYPE.PARAGRAPH)
    risk_probability.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['probability style'].paragraph_format
    paragraph_format.space_after = Pt(0)
    font = risk_probability.style.font
    font.name = font_name
    font.size = Pt(text_font_size)
    probability_title_run = risk_probability.add_run("סבירות מימוש הנזק:")
    probability_title_run.bold = True

    probability_rank_run = risk_probability.add_run(" " + item['Probability'])
    probability_rank_run.bold = True
    font = probability_rank_run.font
    font.color.rgb = RGBColor(*get_color(item['Probability']))

    # Severity
    risk_severity = document.add_paragraph()
    risk_severity.style = document.styles.add_style("severity style", WD_STYLE_TYPE.PARAGRAPH)
    risk_severity.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['severity style'].paragraph_format
    paragraph_format.space_after = Pt(0)
    font = risk_severity.style.font
    font.name = font_name
    font.size = Pt(text_font_size)
    severity_title_run = risk_severity.add_run("חומרת הנזק:")
    severity_title_run.bold = True

    severity_rank_run = risk_severity.add_run(" " + item['Severity'])
    severity_rank_run.bold = True
    font = severity_rank_run.font
    font.color.rgb = RGBColor(*get_color(item['Severity']))


    # Overall
    risk_overall = document.add_paragraph()
    risk_overall.style = document.styles.add_style("overall style", WD_STYLE_TYPE.PARAGRAPH)
    risk_overall.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['overall style'].paragraph_format
    paragraph_format.space_after = Pt(0)
    font = risk_overall.style.font
    font.name = font_name
    font.size = Pt(text_font_size)
    overall_title_run = risk_overall.add_run("חומרת הנזק:")
    overall_title_run.bold = True

    overall_rank_run = risk_overall.add_run(" " + item['OverallRisk'])
    overall_rank_run.bold = True
    font = overall_rank_run.font
    font.color.rgb = RGBColor(*get_color(item['OverallRisk']))

    # Space 3
    space = document.add_paragraph()
    space.style = document.styles.add_style("space3 style", WD_STYLE_TYPE.PARAGRAPH)
    space.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['space3 style'].paragraph_format
    paragraph_format.space_after = Pt(6)
    font = space.style.font
    font.name = font_name
    font.size = Pt(space_font_size)

    # Finding Description Title
    finding_risk_title = document.add_paragraph(":פירוט מימוש הסיכון")
    finding_risk_title.style = document.styles.add_style("finding risk title style", WD_STYLE_TYPE.PARAGRAPH)
    finding_risk_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['finding risk title style'].paragraph_format
    paragraph_format.space_after = Pt(6)
    font = finding_risk_title.style.font
    font.name = font_name
    font.size = Pt(sub_title_font_size)
    font.underline = True

    # Finding Description
    finding_risk_text = reformat_text_dots(item['RiskDetails'])
    finding_risk = document.add_paragraph(finding_risk_text)
    finding_risk.style = document.styles.add_style("finding risk style", WD_STYLE_TYPE.PARAGRAPH)
    finding_risk.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['finding risk style'].paragraph_format
    paragraph_format.space_after = Pt(6)
    font = finding_risk.style.font
    font.name = font_name
    font.size = Pt(text_font_size)

    # Space 4
    space = document.add_paragraph()
    space.style = document.styles.add_style("space4 style", WD_STYLE_TYPE.PARAGRAPH)
    space.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['space4 style'].paragraph_format
    paragraph_format.space_after = Pt(6)
    font = space.style.font
    font.name = font_name
    font.size = Pt(space_font_size)

    # Recommendations Title
    recommendations_title = document.add_paragraph(":המלצות לתיקון")
    recommendations_title.style = document.styles.add_style("reco title style", WD_STYLE_TYPE.PARAGRAPH)
    recommendations_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['reco title style'].paragraph_format
    paragraph_format.space_after = Pt(6)
    font = recommendations_title.style.font
    font.name = font_name
    font.size = Pt(sub_title_font_size)
    font.underline = True

    # Recommendations

    recommendations_list = item['Recommendations'].split("\n")
    formatted_string = ""
    for i in range(len(recommendations_list)):
        r = recommendations_list[i].strip()
        new_r = reformat_recommendation(r)
        formatted_string += new_r + "\n"

    recommendations = document.add_paragraph(formatted_string)
    recommendations.style = document.styles.add_style("reco style", WD_STYLE_TYPE.PARAGRAPH)
    recommendations.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_format = document.styles['reco style'].paragraph_format
    paragraph_format.right_indent = Inches(0.25)
    paragraph_format.space_after = Pt(6)
    font = recommendations.style.font
    font.name = font_name
    font.size = Pt(text_font_size)

    document.save(os.path.join(current_path, dir, item['Uid']) + ".docx")


def get_color(rank_input):
    if rank_input == "נמוכה":
        return green

    elif rank_input == "בינונית":
        return yellow

    elif rank_input == "גבוהה":
        return green


def reformat_text_dots(text):
    # check if '.' in the end of the string
    # if true - reformat the string for rtl purposes
    if text[len(text)-1] == ".":
        sliced_text = text[:len(text)-1]
        # print(sliced_text)
        return "." + sliced_text
    return text


def reformat_recommendation(text):
    # check if '.' in the end of the string
    # if true - reformat the string for rtl purposes
    if text[len(text)-1] == ".":
        text = text[:len(text)-1]
    return "." + text + " ✓"
